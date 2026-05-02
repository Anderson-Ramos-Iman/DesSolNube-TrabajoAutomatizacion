from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Título principal ──
titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run('INFORME TÉCNICO')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitulo.add_run('Sistema Automatizado de Registro de Pagos – CobrApp')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# ── Funciones auxiliares ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_normal(text):
    return doc.add_paragraph(text)

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

# ══════════════════════════════════════════════
# 1. Arquitectura del Sistema
# ══════════════════════════════════════════════
add_heading_styled('1. Arquitectura del Sistema')
add_normal(
    'El sistema CobrApp fue diseñado bajo una arquitectura orientada a microservicios (SOA), '
    'donde cada componente opera de forma independiente y se comunica mediante protocolos HTTP/REST. '
    'Este enfoque permite un alto grado de desacoplamiento entre los módulos, favoreciendo la escalabilidad horizontal, '
    'el mantenimiento evolutivo del código y la operación continua del sistema sin dependencia de infraestructura local.'
)
add_normal(
    'La decisión de distribuir la lógica en servicios independientes responde a los principios de '
    'separation of concerns y single responsibility, pilares fundamentales en el desarrollo de software moderno. '
    'De esta manera, cada servicio se encarga exclusivamente de una tarea específica dentro del pipeline de procesamiento, '
    'lo que reduce el acoplamiento y facilita la depuración ante eventuales fallos en producción.'
)

add_heading_styled('1.1 Componentes principales', level=2)
add_normal('A continuación se describen los componentes clave que conforman la arquitectura del sistema:')
add_bullet(
    'Constituye el punto de entrada del sistema. El usuario final interactúa enviando comprobantes de pago '
    '(capturas de pantalla de Yape, Plin o Lemon) directamente al grupo de Telegram, sin necesidad de instalar '
    'aplicaciones adicionales ni acceder a interfaces web externas.',
    bold_prefix='Cliente (Usuario): '
)
add_bullet(
    'Actúa como la capa de interfaz (frontend layer) del sistema. A través de la Telegram Bot API, '
    'el bot intercepta los mensajes entrantes, valida el tipo de contenido (MIME type de imagen) '
    'y gestiona las respuestas automáticas hacia el usuario mediante mensajes enriquecidos.',
    bold_prefix='Bot de Telegram: '
)
add_bullet(
    'Funciona como el orquestador central (workflow engine) del sistema. Implementa toda la lógica de negocio '
    'mediante un flujo de nodos interconectados que incluyen triggers, condicionales (IF nodes), '
    'HTTP requests, operaciones CRUD sobre Google Sheets y callbacks de notificación. '
    'Su ejecución en la nube garantiza disponibilidad 24/7 sin necesidad de mantener un servidor local activo.',
    bold_prefix='n8n Cloud (Orquestador): '
)
add_bullet(
    'Microservicio desarrollado en Python con el framework FastAPI, desplegado como contenedor Docker en Render. '
    'Se encarga del procesamiento de imágenes mediante técnicas de Computer Vision (OpenCV) y reconocimiento óptico '
    'de caracteres (OCR) con Tesseract. Expone endpoints RESTful que reciben la imagen como multipart/form-data '
    'y retornan un JSON estructurado con los datos extraídos del comprobante.',
    bold_prefix='API OCR (FastAPI en Render): '
)
add_bullet(
    'Cumple la función de capa de persistencia (data layer) del sistema. Fue seleccionado como datastore '
    'por su facilidad de integración nativa con n8n, su accesibilidad en tiempo real y la posibilidad de '
    'realizar consultas de búsqueda (lookup) para la validación de registros duplicados mediante '
    'comparación de campos clave como el número de operación.',
    bold_prefix='Google Sheets: '
)

add_heading_styled('1.2 Flujo de funcionamiento', level=2)
add_normal(
    'El pipeline de procesamiento sigue un flujo secuencial event-driven, donde cada etapa se activa '
    'como consecuencia del resultado de la anterior. A continuación se detalla el flujo completo de ejecución:'
)
flujo_pasos = [
    'El usuario envía una imagen del comprobante de pago al grupo de Telegram, lo cual dispara un evento de tipo message con contenido multimedia.',
    'El Telegram Trigger configurado en n8n detecta el evento entrante y captura el payload del mensaje, incluyendo el file_id del recurso multimedia.',
    'Se ejecuta un nodo condicional (IF node) que evalúa si el mensaje contiene un objeto photo válido, descartando mensajes de texto, stickers u otros tipos de contenido no procesable.',
    'Mediante el nodo Get File de Telegram, se realiza un request a la API de Telegram para descargar el archivo binario de la imagen utilizando el file_id obtenido previamente.',
    'Se envía la imagen al endpoint POST /procesar-imagen de la API desplegada en Render mediante un HTTP Request con content-type multipart/form-data.',
    'La API recibe el archivo, aplica preprocesamiento con OpenCV (conversión a escala de grises, resize y thresholding), ejecuta el motor OCR de Tesseract y aplica expresiones regulares (regex) para extraer los campos relevantes: monto, nombre, número de operación, fecha y hora.',
    'n8n evalúa el campo "valido" del JSON response para determinar si la extracción fue exitosa. En caso de fallo, se notifica al usuario con un mensaje descriptivo del error.',
    'Se ejecuta una consulta lookup en Google Sheets comparando el número de operación extraído contra los registros existentes, aplicando normalización de datos (eliminación de leading zeros) para evitar falsos negativos.',
    'Si no se detecta duplicado, se ejecuta un Append Row en la hoja de cálculo con todos los campos estructurados y se envía un mensaje de confirmación al grupo de Telegram con el resumen del pago registrado.',
    'Si se detecta un registro duplicado, el sistema notifica al usuario indicando que la operación ya fue registrada previamente, evitando la inserción de datos redundantes.',
    'De forma programada mediante un Schedule Trigger, se genera un reporte diario automatizado que consolida el total recaudado, la cantidad de pagos procesados y la fecha correspondiente.',
]
for i, paso in enumerate(flujo_pasos, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(paso)

# ══════════════════════════════════════════════
# 2. Decisiones Técnicas
# ══════════════════════════════════════════════
add_heading_styled('2. Decisiones Técnicas')
add_normal(
    'Cada tecnología fue seleccionada tras un análisis comparativo considerando factores como curva de aprendizaje, '
    'compatibilidad con el ecosistema del proyecto, costo operativo y facilidad de integración. '
    'A continuación se fundamenta cada decisión técnica adoptada durante el desarrollo.'
)

add_heading_styled('2.1 Uso de n8n Cloud', level=2)
add_normal(
    'Se optó por n8n Cloud como workflow automation engine debido a que proporciona un entorno de ejecución '
    'serverless que elimina la necesidad de provisionar infraestructura propia. Sus principales ventajas son:'
)
add_bullet('Ejecución persistente 24/7 en la nube, sin depender de un entorno local o una máquina encendida permanentemente.')
add_bullet('Conectores nativos (built-in nodes) para Telegram Bot API y Google Sheets API, reduciendo significativamente el tiempo de desarrollo.')
add_bullet('Interfaz visual drag-and-drop para la construcción de flujos, lo que facilita la iteración rápida y el debugging visual del pipeline.')
add_bullet('Soporte robusto para lógica condicional mediante IF nodes, switch nodes y expresiones JavaScript inline.')

add_heading_styled('2.2 Uso de FastAPI + OCR', level=2)
add_normal(
    'Para el backend de procesamiento de imágenes se implementó una API RESTful con FastAPI, '
    'un framework asíncrono de alto rendimiento basado en Python y tipado con Pydantic. '
    'La elección se fundamenta en los siguientes criterios:'
)
add_bullet('Alto throughput y baja latencia gracias a su naturaleza asíncrona (async/await) basada en Starlette y Uvicorn como ASGI server.')
add_bullet('Generación automática de documentación interactiva (Swagger UI / OpenAPI) en el endpoint /docs, facilitando el testing y la integración con terceros.')
add_bullet('Ecosistema maduro de librerías de Computer Vision y OCR: pytesseract como wrapper de Tesseract-OCR, OpenCV para preprocesamiento de imágenes (grayscale, thresholding, resize) y NumPy para manipulación eficiente de matrices de píxeles.')
add_normal(
    'El pipeline de OCR implementa un flujo de preprocesamiento → extracción → parsing con regex, '
    'donde cada etapa aplica transformaciones específicas para maximizar la precisión de la lectura. '
    'Se utilizan expresiones regulares compiladas para extraer patrones de monto (S/ o PEN), '
    'número de operación, fecha en múltiples formatos y nombre del titular.'
)

add_heading_styled('2.3 Uso de Docker', level=2)
add_normal(
    'La containerización con Docker fue una decisión estratégica orientada a garantizar la portabilidad '
    'y reproducibilidad del entorno de ejecución. El Dockerfile implementa un build multi-stage sobre '
    'la imagen base python:3.11-slim, e incluye la instalación de dependencias del sistema como '
    'tesseract-ocr, tesseract-ocr-spa y librerías gráficas (libglib2.0, libgl1).'
)
add_bullet('Encapsulamiento completo del runtime environment, eliminando el clásico problema de "funciona en mi máquina".')
add_bullet('Gestión declarativa de dependencias mediante requirements.txt y pip install dentro del contenedor.')
add_bullet('Compatibilidad directa con plataformas PaaS como Render, que soportan despliegue nativo desde Dockerfile.')

add_heading_styled('2.4 Uso de Render', level=2)
add_normal(
    'Render fue seleccionado como plataforma de hosting (PaaS) para el despliegue de la API debido a '
    'su soporte nativo para Docker deployments y su modelo de pricing que incluye un tier gratuito '
    'adecuado para proyectos académicos y prototipos funcionales.'
)
add_bullet('Despliegue automatizado desde repositorio Git con build pipeline integrado basado en Dockerfile.')
add_bullet('Tier gratuito con recursos suficientes para manejar el throughput requerido por el proyecto.')
add_bullet('Provisión automática de URL pública con certificado SSL, accesible desde cualquier servicio externo como n8n Cloud.')

add_heading_styled('2.5 Uso de Google Sheets', level=2)
add_normal(
    'Se adoptó Google Sheets como capa de persistencia del sistema, funcionando como una base de datos '
    'no relacional ligera. Si bien no es una solución de almacenamiento convencional (como PostgreSQL o MongoDB), '
    'resulta altamente efectiva para este caso de uso por las siguientes razones:'
)
add_bullet('Integración nativa con n8n mediante el nodo Google Sheets, soportando operaciones de lectura (lookup), escritura (append) y búsqueda (search).')
add_bullet('Acceso en tiempo real con interfaz visual que permite monitorear los registros sin necesidad de herramientas de administración de base de datos.')
add_bullet('Curva de aprendizaje mínima y configuración zero-setup, sin necesidad de provisionar instancias de base de datos ni gestionar connection strings.')
add_bullet('Capacidad de compartir el documento con stakeholders para auditoría y verificación de datos en tiempo real.')

# ══════════════════════════════════════════════
# 3. Dificultades Encontradas
# ══════════════════════════════════════════════
add_heading_styled('3. Dificultades Encontradas')
add_normal(
    'Durante el proceso de desarrollo e integración se presentaron diversas dificultades técnicas que requirieron '
    'un análisis detallado de root cause y la implementación de soluciones específicas. A continuación se documenta '
    'cada incidencia junto con su diagnóstico y resolución.'
)

dificultades = [
    {
        'titulo': '3.1 Problema con URLs en HTTP Request',
        'error': 'El nodo HTTP Request generaba una URL malformada con doble esquema: http://https://..., provocando un error de conexión (ERR_INVALID_URL).',
        'causa': 'El campo URL del nodo en n8n tenía configurado un prefijo http:// por defecto, y al concatenarse con la URL del servicio que ya incluía https://, se generaba una URI inválida.',
        'solucion': 'Se corrigió el campo URL del nodo removiendo el prefijo redundante, dejando únicamente el esquema correcto: https://dessolnube-trabajoautomatizacion.onrender.com/procesar-imagen.'
    },
    {
        'titulo': '3.2 API en Render en estado "sleep" (cold start)',
        'error': 'Las primeras solicitudes a la API presentaban timeouts o latencias superiores a 30 segundos, causando que n8n abortara la ejecución del workflow.',
        'causa': 'El tier gratuito de Render implementa un mecanismo de cold start que suspende los contenedores inactivos tras 15 minutos sin tráfico. Al recibir una nueva solicitud, el contenedor debe realizar un spin-up completo antes de poder procesar el request.',
        'solucion': 'Se incrementó el timeout del nodo HTTP Request en n8n para tolerar la latencia del cold start. Adicionalmente, se consideró implementar un health-check periódico (cron ping) para mantener el contenedor activo.'
    },
    {
        'titulo': '3.3 Detección incorrecta de duplicados',
        'error': 'El sistema no identificaba correctamente pagos duplicados, permitiendo la inserción de registros con el mismo número de operación.',
        'causa': 'Discrepancia en el formato de datos: el OCR extraía el número de operación con leading zeros (ej: 00286700) mientras que Google Sheets almacenaba el valor sin ellos (286700), causando que la comparación string-to-string retornara false.',
        'solucion': 'Se implementó una normalización de datos aplicando .replace(/^0+/, \'\') en el nodo JavaScript de n8n antes de ejecutar el lookup, garantizando la consistencia en la comparación de identificadores.'
    },
    {
        'titulo': '3.4 Datos no se registraban en Google Sheets',
        'error': 'El flujo se ejecutaba sin errores aparentes, pero no se insertaban nuevas filas en la hoja de cálculo.',
        'causa': 'El nodo Append Row de Google Sheets estaba conectado en una rama incorrecta del flujo condicional, por lo que nunca recibía la ejecución tras la validación.',
        'solucion': 'Se reestructuró la conexión de nodos siguiendo el flujo correcto: IF válido → IF duplicado (rama false) → Append Row → Send Message, asegurando que el registro solo se ejecute cuando los datos son válidos y no duplicados.'
    },
    {
        'titulo': '3.5 Reporte diario mostraba valores en cero',
        'error': 'El reporte generado por el Schedule Trigger mostraba un total recaudado de S/ 0.00 y cantidad de pagos igual a 0, a pesar de existir registros en la hoja.',
        'causa': 'Inconsistencia en el formato de fecha utilizado para filtrar los registros del día actual. El nodo JavaScript comparaba la fecha en formato ISO (2026-05-01) contra el formato almacenado en Sheets (01/05/2026), resultando en cero coincidencias.',
        'solucion': 'Se ajustó la lógica de comparación de fechas en el nodo Code de JavaScript, normalizando ambos valores al formato dd/mm/yyyy antes de ejecutar el filtrado y la sumatoria de montos.'
    },
    {
        'titulo': '3.6 Problema con lectura de archivos en n8n Cloud',
        'error': 'El nodo "Read Binary File" / "Read File from Disk" lanzaba un error de filesystem access al intentar leer archivos locales.',
        'causa': 'n8n Cloud opera en un entorno sandbox sin acceso al filesystem del host, por lo que cualquier operación de I/O sobre archivos locales resulta en un permission denied o file not found.',
        'solucion': 'Se migró la estrategia de almacenamiento, reemplazando la lectura de archivos locales (Excel) por consultas directas a Google Sheets como fuente de datos primaria, eliminando la dependencia del filesystem local.'
    },
]

for d in dificultades:
    add_heading_styled(d['titulo'], level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problema: ')
    run.bold = True
    p.add_run(d['error'])

    p = doc.add_paragraph()
    run = p.add_run('Causa raíz: ')
    run.bold = True
    p.add_run(d['causa'])

    p = doc.add_paragraph()
    run = p.add_run('Solución implementada: ')
    run.bold = True
    p.add_run(d['solucion'])

# ══════════════════════════════════════════════
# 4. Soluciones Implementadas
# ══════════════════════════════════════════════
add_heading_styled('4. Soluciones Implementadas')
add_normal(
    'A lo largo del desarrollo se implementaron múltiples soluciones orientadas a robustecer el sistema, '
    'mejorar la precisión del procesamiento y garantizar la integridad de los datos. Las principales son:'
)
soluciones = [
    'Validación de MIME type en el Telegram Trigger para aceptar exclusivamente mensajes con contenido de tipo imagen (photo), descartando texto, audio, documentos y stickers.',
    'Pipeline de limpieza y sanitización de datos OCR aplicando expresiones regulares para normalizar nombres (remoción de caracteres especiales, letras sueltas) y formatear montos con precisión decimal.',
    'Normalización del número de operación mediante strip de leading zeros para garantizar consistencia en las búsquedas de duplicados.',
    'Control de integridad de datos mediante lookup en Google Sheets antes de cada inserción, evitando registros duplicados que podrían alterar los reportes financieros.',
    'Registro estructurado con campos tipados (Fecha, Hora, Nombre, Monto, Tipo, Operación, Estado, Registrado_por) que facilitan la consulta y el análisis posterior.',
    'Sistema de notificaciones automáticas con mensajes contextuales: confirmación de registro exitoso, alerta de duplicado y notificación de error en la extracción.',
    'Automatización del reporte diario mediante Schedule Trigger con ejecución programada, consolidando métricas clave del día en un mensaje de resumen.',
    'Despliegue completo en servicios cloud (Render + n8n Cloud + Google Sheets) para garantizar disponibilidad continua y operación sin intervención manual.',
]
for s in soluciones:
    add_bullet(s)

# ══════════════════════════════════════════════
# 5. Conclusiones Técnicas
# ══════════════════════════════════════════════
add_heading_styled('5. Conclusiones Técnicas')
add_normal(
    'El desarrollo de CobrApp permitió aplicar en un entorno real los principios de arquitectura distribuida, '
    'integración de servicios en la nube y automatización de procesos. Las principales conclusiones técnicas son:'
)
conclusiones = [
    'La arquitectura basada en microservicios desacoplados demostró ser altamente efectiva para separar responsabilidades (procesamiento OCR, orquestación, persistencia), facilitando el debugging independiente de cada componente y mejorando la escalabilidad del sistema.',
    'El uso de n8n como workflow engine simplifica enormemente la integración entre servicios heterogéneos (Telegram API, FastAPI, Google Sheets API), eliminando la necesidad de desarrollar middleware o backend personalizado para la comunicación entre componentes.',
    'Google Sheets, si bien no es una base de datos relacional convencional, demostró ser una solución pragmática y efectiva como datastore ligero para proyectos de automatización, ofreciendo accesibilidad inmediata, auditoría visual y zero-config deployment.',
    'La containerización con Docker resultó indispensable para garantizar la reproducibilidad del entorno de ejecución y eliminar inconsistencias entre desarrollo local y producción, validando el principio de infrastructure as code.',
    'El motor OCR (Tesseract) es una herramienta poderosa para la extracción de texto desde imágenes, sin embargo, requiere un pipeline robusto de preprocesamiento (thresholding, scaling) y post-procesamiento (regex parsing, data sanitization) para alcanzar niveles de precisión aceptables en producción.',
    'El sistema implementado cumple satisfactoriamente con el objetivo de automatizar el registro de pagos en tiempo real, reduciendo significativamente la carga operativa manual y minimizando el margen de error humano en el proceso de registro.',
]
for c in conclusiones:
    add_bullet(c)

# ══════════════════════════════════════════════
# 6. Estado Final del Sistema
# ══════════════════════════════════════════════
add_heading_styled('6. Estado Final del Sistema')
add_normal(
    'Al cierre del desarrollo, todos los componentes del sistema se encuentran operativos y funcionando '
    'de manera integrada en entornos de producción cloud. El estado de cada módulo es el siguiente:'
)
estados = [
    'Bot de Telegram operativo y respondiendo en tiempo real',
    'API OCR containerizada y desplegada en Render (endpoint activo)',
    'Flujo de automatización en n8n Cloud activo y ejecutándose 24/7',
    'Registro automático de pagos en Google Sheets con campos estructurados',
    'Sistema de detección de duplicados funcional con normalización de datos',
    'Reporte diario automático programado con Schedule Trigger',
]
for e in estados:
    p = doc.add_paragraph()
    run = p.add_run(f'✔ {e}')
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.bold = True

# ── Guardar ──
output_path = r'c:\Users\ander\OneDrive\Documentos\TECSUP\CICLO5-C24A\03_Desarrollo_de_Soluciones_en_la_Nube\TRABAJO01\TrabajoAutomatizacion\Trabajo01-automatizacion\Informe_Tecnico_CobrApp.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
